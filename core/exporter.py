from pathlib import Path


class Exporter:
    @staticmethod
    def to_txt(text, filepath):
        Path(filepath).write_text(text or "", encoding="utf-8")

    @staticmethod
    def to_markdown(text, filepath):
        Path(filepath).write_text(text or "", encoding="utf-8")

    @staticmethod
    def to_srt(segments, filepath):
        if not segments:
            Path(filepath).write_text("", encoding="utf-8")
            return

        lines = []
        for i, seg in enumerate(segments, 1):
            start = Exporter._format_timestamp_srt(seg.get("start", 0))
            end = Exporter._format_timestamp_srt(seg.get("end", 0))
            text = seg.get("text", "")
            lines.append(f"{i}\n{start} --> {end}\n{text}\n")

        Path(filepath).write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def to_vtt(segments, filepath):
        if not segments:
            Path(filepath).write_text("WEBVTT\n\n", encoding="utf-8")
            return

        lines = ["WEBVTT\n"]
        for seg in segments:
            start = Exporter._format_timestamp_vtt(seg.get("start", 0))
            end = Exporter._format_timestamp_vtt(seg.get("end", 0))
            text = seg.get("text", "")
            lines.append(f"{start} --> {end}\n{text}\n")

        Path(filepath).write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def to_docx(text, filepath, title="Transcription"):
        try:
            import docx
        except Exception as exc:
            raise RuntimeError("python-docx is required for DOCX export") from exc

        doc = docx.Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(text or "")
        doc.save(filepath)

    @staticmethod
    def to_pdf(text, filepath, title="Transcription"):
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except Exception as exc:
            raise RuntimeError("reportlab is required for PDF export") from exc

        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, title)
        c.setFont("Helvetica", 11)
        y = height - 80
        for line in (text or "").split("\n"):
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 11)
            c.drawString(50, y, line[:100])
            y -= 15
        c.save()

    @staticmethod
    def _format_timestamp_srt(seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

    @staticmethod
    def _format_timestamp_vtt(seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
